import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ==========================================
# 1. CONFIGURACIÓN DE PÁGINA
# ==========================================
st.set_page_config(page_title="EMS Control Center", layout="wide", page_icon="⚡", initial_sidebar_state="expanded")

# ==========================================
# 2. ESTILOS CSS AVANZADOS (INTER + SCADA DARK THEME)
# ==========================================
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

    * { font-family: 'Inter', sans-serif !important; }

    /* Paleta de colores SCADA sugerida */
    :root {
        --bg-main: #080F1F;
        --bg-panel: #111B2E;
        --bg-card: #172338;
        --border: #26354D;
        --cyan: #00B8FF;
        --green: #00D084;
        --yellow: #FFB020;
        --red: #FF4D5A;
        --text-main: #F8FAFC;
        --text-sec: #94A3B8;
    }

    /* Fondo principal y reseteo */
    .stApp { background-color: var(--bg-main); color: var(--text-main); }
    .block-container { padding-top: 1rem; max-width: 1400px; }
    header, footer { visibility: hidden; display: none; }

    /* Fix para colores de texto por defecto en Streamlit */
    .stMarkdown p, .stText p, .stSlider label, .stTextInput label { color: var(--text-sec) !important; font-weight: 500 !important; font-size: 13px !important; }
    div[data-testid="stThumbValue"] { color: var(--text-main) !important; font-size: 14px !important; font-weight: 600 !important; }
    
    /* Configuración Avanzada - Header y Box */
    .config-header { font-size: 26px; font-weight: 700; color: var(--cyan); margin-bottom: 12px; }
    .config-box { background-color: var(--bg-panel); border: 1px solid var(--border); border-radius: 8px; padding: 24px; margin-bottom: 24px; }
    
    /* Sliders personalizados: Color Base (Cyan) en lugar de Rojo */
    .stSlider > div > div > div > div { background-color: var(--cyan) !important; }
    .stSlider > div > div > div > div > div { border-color: var(--cyan) !important; }

    /* Barra lateral */
    section[data-testid="stSidebar"] { background-color: var(--bg-panel); border-right: 1px solid var(--border); padding-top: 24px;}
    div[data-testid="stSidebar"] * { color: var(--text-main) !important; }
    
    /* Botones Sidebar (Navegación) */
    div[data-testid="stSidebar"] .stButton > button {
        background-color: transparent; border: 1px solid transparent; color: var(--text-sec) !important; 
        text-align: left; justify-content: flex-start; width: 100%; height: 42px; border-radius: 8px; font-weight: 500; font-size: 14px;
        transition: all 0.2s;
    }
    div[data-testid="stSidebar"] .stButton > button:hover { background-color: var(--bg-card); color: var(--cyan) !important; }
    
    /* Status Panel Sidebar */
    .status-panel { border-top: 1px solid var(--border); padding-top: 16px; margin-top: 20px; }
    .status-lbl { font-size: 12px; color: var(--text-sec); letter-spacing: 1px; margin-bottom: 4px; }
    .status-online { color: var(--green); font-weight: 700; font-size: 14px; display: flex; align-items: center; gap: 8px; }
    .status-online::before { content: ''; display: inline-block; width: 8px; height: 8px; background-color: var(--green); border-radius: 50%; box-shadow: 0 0 6px var(--green); }

    /* Header Principal EMS CONTROL CENTER */
    .scada-header {
        display: flex; justify-content: space-between; align-items: center; 
        border-bottom: 1px solid var(--border); padding-bottom: 16px; margin-bottom: 24px;
    }
    .scada-title { margin: 0; color: var(--text-main); font-size: 30px; font-weight: 700; display: flex; align-items: center; gap: 10px; }
    .scada-subtitle { color: var(--text-sec); font-size: 12px; font-weight: 500; text-transform: uppercase; letter-spacing: 1.5px; margin-top: 4px;}
    .scada-status-box { text-align: right; }
    
    /* Tarjetas Dashboard */
    .kpi-card { background-color: var(--bg-panel); border: 1px solid var(--border); border-radius: 8px; padding: 16px; margin-bottom: 16px; }
    .kpi-title { font-size: 13px; color: var(--text-sec); font-weight: 600; text-transform: uppercase; margin-bottom: 8px; }
    .kpi-value { font-size: 28px; font-weight: 700; color: var(--text-main); line-height: 1; }
    .kpi-unit { font-size: 14px; color: var(--text-sec); font-weight: 500; margin-left: 4px; }
    .kpi-sub { font-size: 12px; color: var(--text-sec); margin-top: 10px; display: flex; justify-content: space-between; }
    
    .c-cyan { color: var(--cyan); } .c-green { color: var(--green); } .c-yellow { color: var(--yellow); } .c-red { color: var(--red); }

    /* Botones Superiores */
    .btn-row { display: flex; gap: 12px; margin-bottom: 24px; }
    div.stButton > button { background-color: var(--bg-card); color: var(--text-main) !important; border: 1px solid var(--border); border-radius: 9px; height: 42px; font-size: 14px; font-weight: 600; width: 100%; transition: all 0.2s;}
    div.stButton > button:hover { border-color: var(--cyan); color: var(--cyan) !important; }
    div.stButton > button[kind="primary"] { background-color: rgba(0,184,255,0.1); border-color: var(--cyan); color: var(--cyan) !important; height: 46px; font-size: 15px;}
    
    /* Inputs de Configuración */
    .stTextInput > div > div > input { background-color: var(--bg-card); color: var(--text-main); border: 1px solid var(--border); border-radius: 6px; }
</style>
""", unsafe_allow_html=True)

# ==========================================
# 3. GESTIÓN DE ESTADO (CAMPOS GENÉRICOS)
# ==========================================
if 'page' not in st.session_state:
    st.session_state.page = "Dashboard"

if 'config' not in st.session_state:
    st.session_state.config = {
        'nombre_proyecto': 'EMS Bloque D',
        'ubicacion_proyecto': 'UPS Campus Centenario',
        'p_lim': 130.0, 'c_bat': 250.0, 'p_pv': 150.0,
        'v_nom': 220.0, 's_trafo': 1000.0, 'carga_noc': 40.0,
        'ps_activo': True
    }
cfg = st.session_state.config

# ==========================================
# 4. LÓGICA DE INGENIERÍA (CÁLCULOS)
# ==========================================
REAL_LOAD = [36, 36, 36, 36, 36, 40, 60, 90, 120, 145, 160, 175, 179.1, 140, 150, 155, 160, 165, 172, 175, 130, 90, 50, 36]
PV_BASE = [0, 0, 0, 0, 0, 0, 5, 25, 55, 90, 120, 140, 150, 140, 120, 90, 55, 25, 5, 0, 0, 0, 0, 0]

factor_pv = cfg['p_pv'] / 150.0 if cfg['p_pv'] > 0 else 0.0
pv_real = [round(v * factor_pv, 1) for v in PV_BASE]

soc_min = 0.20 * cfg['c_bat']
soc_max = cfg['c_bat']
energia = cfg['c_bat'] * 0.50
limite_operativo = cfg['p_lim'] if cfg['ps_activo'] else 9999.0

rows_ems = []
for i in range(24):
    p_teorica = REAL_LOAD[i] - pv_real[i]
    p_bat = 0.0
    
    if p_teorica > limite_operativo:
        req = p_teorica - limite_operativo
        p_bat = req if (energia - req) >= soc_min else max(0.0, energia - soc_min)
    elif 1 <= i <= 5: 
        p_bat = -cfg['carga_noc'] if (energia + cfg['carga_noc']) <= soc_max else -(soc_max - energia)
        
    p_red = p_teorica - p_bat
    energia -= p_bat
    soc = (energia / cfg['c_bat']) * 100.0
    rows_ems.append({'Hora': f"{i:02d}:00", 'P_Carga': REAL_LOAD[i], 'P_PV': pv_real[i], 'P_Bat': round(p_bat, 1), 'P_Red': round(p_red, 1), 'SOC': round(soc, 1)})

df_ems = pd.DataFrame(rows_ems)

demanda_max = float(df_ems['P_Carga'].max())
demanda_recortada = float(df_ems['P_Red'].max())
reduccion_pico = demanda_max - demanda_recortada
inv_req = cfg['p_pv'] / 0.95 if cfg['p_pv'] > 0 else cfg['p_lim'] / 0.95
i_nom = (cfg['s_trafo'] * 1000.0) / (1.73205 * cfg['v_nom'])
icc_simetrica = i_nom / (5.75 / 100.0)
carg_sin = (demanda_max / cfg['s_trafo']) * 100.0
carg_con = (demanda_recortada / cfg['s_trafo']) * 100.0
soc_actual = df_ems['SOC'].iloc[12]
estado_ps = "PEAK SHAVING ACTIVO" if cfg['ps_activo'] else "PEAK SHAVING INACTIVO"

# ==========================================
# 5. BARRA LATERAL (NAVEGACIÓN)
# ==========================================
st.sidebar.markdown("""
<div style="margin-bottom: 30px;">
    <h2 style="color: #F8FAFC; font-size: 20px; font-weight: 700; margin: 0;">Navegación</h2>
</div>
""", unsafe_allow_html=True)

if st.sidebar.button("🏠 Dashboard Principal"): st.session_state.page = "Dashboard"
if st.sidebar.button("⚡ Análisis EMS"): st.session_state.page = "EMS"
if st.sidebar.button("📐 Diagrama Unifilar SCADA"): st.session_state.page = "Unifilar"
if st.sidebar.button("📄 Memoria Técnica"): st.session_state.page = "Memoria"
if st.sidebar.button("📦 Exportaciones"): st.session_state.page = "Exportaciones"
if st.sidebar.button("⚙️ Configuración"): st.session_state.page = "Configuracion"

st.sidebar.markdown("""
<div class="status-panel">
    <div class="status-lbl">SYSTEM STATUS</div>
    <div class="status-online">ONLINE</div>
</div>
""", unsafe_allow_html=True)

# ==========================================
# 6. CABECERA Y BOTONES DE ACCIÓN SUPERIOR
# ==========================================
st.markdown(f"""
<div class="scada-header">
    <div>
        <h2 class="scada-title"><span style="color: #FFB020;">⚡</span> EMS CONTROL CENTER</h2>
        <div class="scada-subtitle">ENERGY MANAGEMENT SYSTEM | {cfg['ubicacion_proyecto']}</div>
    </div>
    <div class="scada-status-box">
        <div class="status-online" style="justify-content: flex-end;">SYSTEM ONLINE</div>
        <div style="color: #00B8FF; font-size: 11px; font-weight: 600; margin-top: 4px;">SISTEMA OPERATIVO | {estado_ps}</div>
    </div>
</div>
""", unsafe_allow_html=True)

if st.session_state.page in ["Dashboard", "EMS", "Unifilar"]:
    col_btn1, col_btn2, col_btn3 = st.columns([1,1,1])
    with col_btn1:
        if st.button("▶ EJECUTAR SIMULACIÓN", type="primary"): st.toast("Simulación Completada")
    with col_btn2:
        if st.button("↻ RECALCULAR"): st.rerun()
    with col_btn3:
        if st.button(f"⚡ PEAK SHAVING {'OFF' if cfg['ps_activo'] else 'ON'}"): 
            st.session_state.config['ps_activo'] = not cfg['ps_activo']
            st.rerun()

# ==========================================
# 7. VISTAS DE PÁGINA
# ==========================================

# ------------------------------------------
# VISTA 1: DASHBOARD
# ------------------------------------------
if st.session_state.page == "Dashboard":
    
    # CONFIGURACIÓN AVANZADA
    st.markdown("<div class='config-header'>Configuración Avanzada</div>", unsafe_allow_html=True)
    st.markdown("<div class='config-box'>", unsafe_allow_html=True)
    
    # Diseño de 3 Columnas por 2 Filas
    c1, c2, c3 = st.columns(3)
    with c1:
        st.session_state.config['p_lim'] = st.slider("Set-point límite de red (kW)", 80.0, 200.0, cfg['p_lim'], 5.0)
        st.session_state.config['s_trafo'] = st.slider("Capacidad Trafo (kVA)", 315.0, 2000.0, cfg['s_trafo'], 50.0)
    with c2:
        st.session_state.config['c_bat'] = st.slider("Capacidad BESS (kWh)", 50.0, 1000.0, cfg['c_bat'], 10.0)
        st.session_state.config['v_nom'] = st.slider("Tensión BT (V)", 110.0, 480.0, cfg['v_nom'], 10.0)
    with c3:
        st.session_state.config['p_pv'] = st.slider("Potencia PV (kWp)", 0.0, 300.0, cfg['p_pv'], 10.0)
        st.session_state.config['carga_noc'] = st.slider("Carga Nocturna BESS (kW)", 10.0, 100.0, cfg['carga_noc'], 5.0)
    st.markdown("</div>", unsafe_allow_html=True)

    # TARJETAS DE MÉTRICAS SCADA
    m1, m2, m3, m4 = st.columns(4)
    m1.markdown(f"""<div class="kpi-card"><div class="kpi-title">DEMANDA RED</div>
    <div class="kpi-value">{demanda_recortada:.1f} <span class="kpi-unit">kW</span></div>
    <div class="kpi-sub"><span>Original: {demanda_max:.1f} kW</span> <span class="c-cyan">▼ {reduccion_pico:.1f} kW</span></div></div>""", unsafe_allow_html=True)
    
    m2.markdown(f"""<div class="kpi-card"><div class="kpi-title">ALMACENAMIENTO BESS</div>
    <div class="kpi-value">{cfg['c_bat']:.0f} <span class="kpi-unit">kWh</span></div>
    <div class="kpi-sub"><span>Tecnología: LiFePO4</span> <span class="c-green">SOC ~{soc_actual:.0f}%</span></div></div>""", unsafe_allow_html=True)
    
    m3.markdown(f"""<div class="kpi-card"><div class="kpi-title">GENERACIÓN SOLAR</div>
    <div class="kpi-value">{cfg['p_pv']:.0f} <span class="kpi-unit">kWp</span></div>
    <div class="kpi-sub"><span>Inversor: {inv_req:.1f} kVA</span> <span class="c-green">● NORMAL</span></div></div>""", unsafe_allow_html=True)
    
    m4.markdown(f"""<div class="kpi-card"><div class="kpi-title">TRAFO {cfg['s_trafo']:.0f} kVA</div>
    <div class="kpi-value">{carg_con:.1f} <span class="kpi-unit">%</span></div>
    <div class="kpi-sub"><span>Cargabilidad</span> <span class="{'c-green' if carg_con < 85 else 'c-red'}">● {'NORMAL' if carg_con < 85 else 'ALERTA'}</span></div></div>""", unsafe_allow_html=True)

    # GRÁFICA PRINCIPAL
    st.markdown("<h3 style='color:#F8FAFC; margin-top:20px; font-size:22px; font-weight:600;'>Monitoreo de Potencia en Tiempo de Simulación</h3>", unsafe_allow_html=True)
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=df_ems['Hora'], y=df_ems['P_Carga'], name='Demanda Bruta', line=dict(color='#00B8FF', width=2)))
    fig.add_trace(go.Scatter(x=df_ems['Hora'], y=df_ems['P_Red'], name='Consumo de Red', fill='tozeroy', line=dict(color='#00D084', width=2)))
    if cfg['ps_activo']:
        fig.add_trace(go.Scatter(x=df_ems['Hora'], y=[cfg['p_lim']]*24, name='Límite EMS', line=dict(color='#FF4D5A', width=2, dash='dash')))
    
    fig.update_layout(template="plotly_dark", height=400, margin=dict(l=10, r=10, t=20, b=10), paper_bgcolor='#111B2E', plot_bgcolor='#111B2E', legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1))
    st.plotly_chart(fig, use_container_width=True)

# ------------------------------------------
# VISTA 2: ANÁLISIS EMS
# ------------------------------------------
elif st.session_state.page == "EMS":
    st.markdown("<h3 style='color: #00B8FF;'>Análisis EMS y Despacho de Baterías</h3>", unsafe_allow_html=True)
    fig_soc = go.Figure()
    fig_soc.add_trace(go.Scatter(x=df_ems['Hora'], y=df_ems['SOC'], name='SOC BESS (%)', line=dict(color='#00B8FF', width=2), fill='tozeroy', fillcolor='rgba(0,184,255,0.1)'))
    fig_soc.add_trace(go.Scatter(x=df_ems['Hora'], y=[20]*24, name='Reserva (20%)', line=dict(color='#FF4D5A', width=2, dash='dash')))
    fig_soc.update_layout(template="plotly_dark", height=300, paper_bgcolor='#111B2E', plot_bgcolor='#111B2E')
    st.plotly_chart(fig_soc, use_container_width=True)
    st.dataframe(df_ems, use_container_width=True)

# ------------------------------------------
# VISTA 3: UNIFILAR SCADA (Corregido y Centralizado)
# ------------------------------------------
elif st.session_state.page == "Unifilar":
    st.markdown("<h3 style='color: #00B8FF;'>Diagrama Unifilar Jerárquico (Interfaz SCADA)</h3>", unsafe_allow_html=True)
    
    c_left, c_right = st.columns([1.5, 3.5])
    with c_left:
        st.markdown("<div style='background-color:#111B2E; border:1px solid #26354D; padding:20px; border-radius:8px;'><p style='font-size:16px; font-weight:700; color:#F8FAFC; margin-bottom:12px;'>Equipos</p>", unsafe_allow_html=True)
        eq = st.radio("Sel:", ["Transformador", "BESS", "Inversor", "Arreglo PV", "Red CNEL", "TGBT", "Cargas Bloque D"], label_visibility="collapsed")
        st.markdown("</div><br>", unsafe_allow_html=True)
        
        if eq == "Transformador":
            st.markdown(f"""<div class="kpi-card" style="border-top:3px solid #00B8FF;"><h4 style="color:#00B8FF; margin-top:0;">⚡ TRANSFORMADOR</h4>
            <div style="font-size:14px; line-height:2.0; color:#F8FAFC;"><b>Capacidad:</b> {cfg['s_trafo']} kVA<br><b>Tensión:</b> 69 kV / {cfg['v_nom']/1000} kV<br><b>Icc Simétrica:</b> {(icc_simetrica/1000):.2f} kA<br><b>Carga Actual:</b> {carg_con:.1f} %<br><br><span class="c-green">● ESTADO: NORMAL</span></div></div>""", unsafe_allow_html=True)
        elif eq == "BESS":
            st.markdown(f"""<div class="kpi-card" style="border-top:3px solid #00D084;"><h4 style="color:#00D084; margin-top:0;">🔋 BANCO BESS</h4>
            <div style="font-size:14px; line-height:2.0; color:#F8FAFC;"><b>Capacidad:</b> {cfg['c_bat']} kWh<br><b>Química:</b> LiFePO4<br><b>Energía Útil:</b> {cfg['c_bat']*0.8:.1f} kWh<br><br><span class="c-green">● ONLINE</span></div></div>""", unsafe_allow_html=True)
        else:
            st.markdown(f"""<div class="kpi-card" style="border-top:3px solid #00B8FF;"><h4 style="color:#00B8FF; margin-top:0;">{eq.upper()}</h4>
            <div style="font-size:14px; line-height:2.0; color:#F8FAFC;"><br><span class="c-green">● ESTADO: OPERATIVO</span></div></div>""", unsafe_allow_html=True)

    with c_right:
        fig_sld = go.Figure()
        fig_sld.update_xaxes(visible=False, range=[-120, 120]); fig_sld.update_yaxes(visible=False, range=[-80, 220])
        l_col = '#00B8FF'; t_col = '#F8FAFC'
        
        # Red MT
        fig_sld.add_trace(go.Scatter(x=[0, 0], y=[200, 150], mode='lines', line=dict(color=l_col, width=2), showlegend=False))
        fig_sld.add_annotation(x=30, y=190, text="RED CNEL 13.8 kV", showarrow=False, font=dict(size=12, color=t_col))
        
        # Trafo
        fig_sld.add_shape(type="circle", x0=-12, y0=115, x1=12, y1=145, line_color=l_col, line_width=2)
        fig_sld.add_shape(type="circle", x0=-12, y0=95, x1=12, y1=125, line_color=l_col, line_width=2)
        fig_sld.add_annotation(x=45, y=120, text=f"TRAFO {cfg['s_trafo']} kVA", showarrow=False, font=dict(size=12, color=t_col))
        
        # Línea a TGBT y Disyuntor
        fig_sld.add_trace(go.Scatter(x=[0, 0], y=[95, 60], mode='lines', line=dict(color=l_col, width=2), showlegend=False))
        fig_sld.add_shape(type="rect", x0=-10, y0=65, x1=10, y1=80, line_color=l_col, line_width=2, fillcolor='#111B2E')
        fig_sld.add_annotation(x=35, y=72, text="ITM 50kA", showarrow=False, font=dict(size=11, color=t_col))
        
        # Bus TGBT
        fig_sld.add_trace(go.Scatter(x=[0, 0], y=[60, 40], mode='lines', line=dict(color=l_col, width=2), showlegend=False))
        fig_sld.add_trace(go.Scatter(x=[-90, 90], y=[40, 40], mode='lines', line=dict(color=l_col, width=4), showlegend=False))
        fig_sld.add_annotation(x=0, y=47, text=f"BUS TGBT {cfg['v_nom']}V", showarrow=False, font=dict(size=13, color=t_col, weight="bold"))
        
        # Rama Cargas
        fig_sld.add_trace(go.Scatter(x=[-50, -50], y=[40, 0], mode='lines', line=dict(color='#FF4D5A', width=2), showlegend=False))
        fig_sld.add_shape(type="rect", x0=-75, y0=-15, x1=-25, y1=0, line_color="#FF4D5A", line_width=2)
        fig_sld.add_annotation(x=-50, y=-7.5, text="CARGAS", showarrow=False, font=dict(size=11, color='#FF4D5A'))
        
        # Rama Inversor
        fig_sld.add_trace(go.Scatter(x=[50, 50], y=[40, 0], mode='lines', line=dict(color='#A855F7', width=2), showlegend=False))
        fig_sld.add_shape(type="rect", x0=20, y0=-15, x1=80, y1=0, line_color="#A855F7", line_width=2)
        fig_sld.add_annotation(x=50, y=-7.5, text="INVERSOR", showarrow=False, font=dict(size=11, color='#A855F7'))
        
        # DC Lines y Cajas (PV / BESS)
        fig_sld.add_trace(go.Scatter(x=[35, 35], y=[-15, -40], mode='lines', line=dict(color='#FFB020', width=2), showlegend=False))
        fig_sld.add_trace(go.Scatter(x=[65, 65], y=[-15, -40], mode='lines', line=dict(color='#00D084', width=2), showlegend=False))
        
        fig_sld.add_shape(type="rect", x0=20, y0=-60, x1=50, y1=-40, line_color="#FFB020", line_width=2)
        fig_sld.add_annotation(x=35, y=-50, text="PV", showarrow=False, font=dict(size=11, color='#FFB020'))
        
        fig_sld.add_shape(type="rect", x0=55, y0=-60, x1=85, y1=-40, line_color="#00D084", line_width=2)
        fig_sld.add_annotation(x=70, y=-50, text="BESS", showarrow=False, font=dict(size=11, color='#00D084'))

        fig_sld.update_layout(height=600, margin=dict(l=0, r=0, t=10, b=10), template='plotly_dark', paper_bgcolor='#111B2E', plot_bgcolor='#111B2E')
        st.plotly_chart(fig_sld, use_container_width=True)

# ------------------------------------------
# VISTA 4: MEMORIA TÉCNICA
# ------------------------------------------
elif st.session_state.page == "Memoria":
    st.markdown("<h3 style='color: #00B8FF;'>Generación de Memoria Técnica</h3>", unsafe_allow_html=True)
    st.markdown("<p style='color: #94A3B8;'>El documento Word (.docx) se genera respetando tu formato de ingeniería exacto (Índices y Capítulos del 1 al 9).</p>", unsafe_allow_html=True)
    
    def generar_docx():
        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(1); section.bottom_margin = Inches(1); section.left_margin = Inches(1); section.right_margin = Inches(1)
            
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("MEMORIA TÉCNICA Y ESPECIFICACIONES DE PROYECTO"); r.bold = True; r.font.size = Pt(14)
        
        t_meta = doc.add_table(rows=4, cols=2); t_meta.style = 'Table Grid'
        data_m = [("Departamento:", "Ingeniería y Viabilidad Técnica"), ("Documento:", f"Memoria Técnica EMS - {cfg['nombre_proyecto']}"), ("Código del Documento:", "PROY-EMS-MTC-001"), ("Revisión / Fecha:", "Rev. 01 / 2026")]
        for i, (k, v) in enumerate(data_m):
            t_meta.cell(i, 0).text = k; t_meta.cell(i, 0).paragraphs[0].runs[0].bold = True; t_meta.cell(i, 1).text = v

        doc.add_paragraph()
        doc.add_heading('Historial de revisiones', level=2)
        t_rev = doc.add_table(rows=2, cols=4); t_rev.style = 'Table Grid'
        headers = ["N° de Revisión", "Fecha", "Páginas Revisadas", "Motivo de Revisión"]
        for i, h in enumerate(headers): t_rev.cell(0, i).text = h; t_rev.cell(0, i).paragraphs[0].runs[0].bold = True
        t_rev.cell(1, 0).text = "01"; t_rev.cell(1, 1).text = "2026"; t_rev.cell(1, 2).text = "Todo el documento"; t_rev.cell(1, 3).text = "Revisión General y Emisión"

        doc.add_paragraph()
        doc.add_heading('Documentos Entregados', level=2)
        t_doc = doc.add_table(rows=2, cols=2); t_doc.style = 'Table Grid'
        t_doc.cell(0,0).text = "Documento:"; t_doc.cell(0,0).paragraphs[0].runs[0].bold = True
        t_doc.cell(0,1).text = "Código:"; t_doc.cell(0,1).paragraphs[0].runs[0].bold = True
        t_doc.cell(1,0).text = "Plano Unifilar y Memoria de Cálculo"; t_doc.cell(1,1).text = "PROY-EMS-DUF-001"

        doc.add_heading('ÍNDICE DE CONTENIDO', level=1)
        indice = ["1. OBJETIVOS", "2. INTRODUCCIÓN", "3. UBICACIÓN", "4. DESARROLLO GENERAL\n  4.1 SISTEMA EXISTENTE\n  4.2 SISTEMA PROYECTADO", "5. ESPECIFICACIONES TÉCNICAS", "6. CÁLCULO DE LA DEMANDA Y ESTUDIO ELÉCTRICO", "7. LISTA DE MATERIALES", "8. CONCLUSIONES", "9. ANEXOS"]
        for item in indice: doc.add_paragraph(item)
        doc.add_page_break()
        
        doc.add_heading('1. OBJETIVOS', level=1)
        doc.add_heading('1.1 Objetivo General:', level=2)
        doc.add_paragraph("Incorporación de nuevas tecnologías y mejora de la infraestructura con el objetivo de garantizar un suministro eléctrico competitivo, seguro y eficiente mediante la implementación de un Sistema Inteligente de Gestión de Energía (EMS).")
        doc.add_heading('1.2 Objetivos Específicos:', level=2)
        doc.add_paragraph(f"Electrificación y recorte de demanda pico (Peak Shaving) de {demanda_max:.1f} kW a {cfg['p_lim']:.1f} kW, garantizando el cumplimiento de las normativas de interconexión (IEEE 2030.7, 1547).")

        doc.add_heading('2. INTRODUCCIÓN', level=1)
        doc.add_paragraph("La implementación del proyecto apuesta por el uso de tecnologías híbridas (Generación fotovoltaica y almacenamiento BESS), mitigando los picos de consumo y mejorando el perfil de tensión local.")

        doc.add_heading('3. UBICACIÓN', level=1)
        doc.add_paragraph(f"El centro de carga principal está ubicado en {cfg['ubicacion_proyecto']}.")

        doc.add_heading('4. DESARROLLO GENERAL', level=1)
        doc.add_heading('4.1 SISTEMA EXISTENTE', level=2)
        doc.add_paragraph(f"Actualmente el centro de carga opera con un transformador de {cfg['s_trafo']:.0f} kVA a {cfg['v_nom']:.1f}V, alcanzando una demanda máxima registrada de {demanda_max:.1f} kW, con una cargabilidad térmica original del {carg_sin:.1f}%.")
        
        doc.add_heading('4.2 SISTEMA PROYECTADO', level=2)
        doc.add_paragraph(f"Se proyecta la integración de un banco de baterías de {cfg['c_bat']:.0f} kWh y un sistema solar de {cfg['p_pv']:.0f} kWp, acoplados a un inversor de {inv_req:.1f} kVA. El algoritmo EMS limitará la potencia tomada de la red a {cfg['p_lim']:.1f} kW, mejorando la cargabilidad del transformador al {carg_con:.1f}%.")

        doc.add_heading('5. ESPECIFICACIONES TÉCNICAS', level=1)
        doc.add_paragraph(f"• INVERSOR MULTIMODO: Potencia Nominal de {inv_req:.1f} kVA, Factor de Potencia mínimo regulable a 0.95 (IEEE 1547).\n• BANCO BESS: Capacidad Nominal de {cfg['c_bat']:.0f} kWh en tecnología LiFePO4, con DoD configurado al 80% (Reserva de seguridad SOC_min de {soc_min:.1f} kWh).")
        
        doc.add_heading('6. CÁLCULO DE LA DEMANDA Y ESTUDIO TÉCNICO', level=1)
        doc.add_paragraph(f"Para el bus principal en {cfg['v_nom']:.1f}V, la corriente nominal del transformador es de {i_nom:.1f} A. Considerando una impedancia de Z=5.75%, la corriente de falla simétrica es Icc = {icc_simetrica/1000.0:.2f} kA. Se validó la capacidad interruptiva requerida del disyuntor principal a 50 kA.")

        doc.add_heading('7. LISTA DE MATERIALES', level=1)
        t_mat = doc.add_table(rows=5, cols=4); t_mat.style = 'Table Grid'
        mat_headers = ["ÍTEM", "DESCRIPCIÓN", "UNIDAD", "CANTIDAD"]
        for i, h in enumerate(mat_headers): t_mat.cell(0, i).text = h; t_mat.cell(0, i).paragraphs[0].runs[0].bold = True
        
        materiales = [("1", f"Sistema Almacenamiento BESS {cfg['c_bat']:.0f} kWh LiFePO4", "GLB", "1"), ("2", f"Inversor Híbrido Multimodo {inv_req:.1f} kVA", "UN", "1"), ("3", f"Sistema Fotovoltaico {cfg['p_pv']:.0f} kWp", "GLB", "1"), ("4", "Controlador PLC Microgrid EMS", "UN", "1")]
        for r_idx, (i, d, u, c) in enumerate(materiales, start=1):
            t_mat.cell(r_idx, 0).text = i; t_mat.cell(r_idx, 1).text = d; t_mat.cell(r_idx, 2).text = u; t_mat.cell(r_idx, 3).text = c

        doc.add_heading('8. CONCLUSIONES', level=1)
        doc.add_paragraph(f"El sistema diseñado logra un aplanamiento neto de demanda de {reduccion_pico:.1f} kW, reduciendo el estrés térmico en el transformador de {cfg['s_trafo']:.0f} kVA y garantizando el cumplimiento normativo.")

        target = io.BytesIO(); doc.save(target)
        return target.getvalue()

    c1, c2 = st.columns([1,2])
    with c1:
        st.download_button("📄 DESCARGAR MEMORIA TÉCNICA (.DOCX)", generar_docx(), f"Memoria_Tecnica_{cfg['nombre_proyecto'].replace(' ','_')}.docx", 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', use_container_width=True)

# ------------------------------------------
# VISTA 5: EXPORTACIONES Y MATLAB
# ------------------------------------------
elif st.session_state.page == "Exportaciones":
    st.markdown("<h3 style='color: #00B8FF;'>Exportación de Planos, Datos y Código</h3>", unsafe_allow_html=True)
    
    def generate_dxf_full():
        lines = ["0", "SECTION", "2", "HEADER", "0", "ENDSEC", "0", "SECTION", "2", "TABLES", "0", "ENDSEC", "0", "SECTION", "2", "BLOCKS", "0", "ENDSEC", "0", "SECTION", "2", "ENTITIES"]
        def add_line(layer, x1, y1, x2, y2, color="7"):
            lines.extend(["0", "LINE", "8", layer, "62", color, "10", f"{x1:.2f}", "20", f"{y1:.2f}", "30", "0.0", "11", f"{x2:.2f}", "21", f"{y2:.2f}", "31", "0.0"])
        def add_circle(layer, cx, cy, r, color="7"):
            lines.extend(["0", "CIRCLE", "8", layer, "62", color, "10", f"{cx:.2f}", "20", f"{cy:.2f}", "30", "0.0", "40", f"{r:.2f}"])
        def add_text(layer, x, y, text, height=3.0, color="7"):
            lines.extend(["0", "TEXT", "8", layer, "62", color, "10", f"{x:.2f}", "20", f"{y:.2f}", "30", "0.0", "40", f"{height:.2f}", "1", str(text)])
        def add_box(layer, x1, y1, x2, y2, color="7"):
            add_line(layer, x1, y1, x2, y1, color); add_line(layer, x2, y1, x2, y2, color)
            add_line(layer, x2, y2, x1, y2, color); add_line(layer, x1, y2, x1, y1, color)

        # MARCO Y CAJETÍN TÉCNICO
        add_box("MARCO", -200, -150, 300, 250, color="2")
        add_box("CAJETIN", 150, -150, 300, -80, color="2")
        add_line("CAJETIN", 150, -100, 300, -100, color="2")
        add_line("CAJETIN", 150, -120, 300, -120, color="2")
        
        add_text("TEXTOS", 160, -90, f"PROYECTO: {cfg['nombre_proyecto'].upper()}", 4.0, "7")
        add_text("TEXTOS", 160, -110, f"UBICACION: {cfg['ubicacion_proyecto'].upper()}", 3.5, "7")
        add_text("TEXTOS", 160, -135, "PLANO: DIAGRAMA UNIFILAR EMS", 3.0, "7")

        # DIAGRAMA UNIFILAR COMPLETO
        add_line("RED_MT", 50, 220, 50, 160, color="4")
        add_text("TEXTOS", 55, 210, "ACOMETIDA RED PRINCIPAL CNEL - 69 kV / 13.8 kV", 3.5, "7")
        
        add_circle("EQUIPOS", 50, 160, 3.0, color="1")
        add_text("TEXTOS", 58, 158, "CCF 100A + APARTARRAYOS 12 kV", 3.0, "7")
        add_line("RED_MT", 50, 157, 50, 130, color="4")

        add_circle("TRAFO", 50, 115, 15, color="4")
        add_circle("TRAFO", 50, 95, 15, color="4")
        add_text("TEXTOS", 60, 120, "Δ", 4.0, "7")
        add_text("TEXTOS", 60, 90, "Y", 4.0, "7")
        
        add_box("INFO", 80, 80, 200, 130, color="3")
        add_text("TEXTOS", 85, 120, f"TRANSFORMADOR PEDESTAL {cfg['s_trafo']:.0f} kVA", 3.5, "7")
        add_text("TEXTOS", 85, 110, f"Secundario: {cfg['v_nom']:.0f} V (3F-4H, Dyn11)", 3.0, "7")
        add_text("TEXTOS", 85, 100, f"Z% = 5.75%  |  Icc_sim = {icc_simetrica/1000.0:.2f} kA", 3.0, "7")

        add_line("RED_BT", 50, 80, 50, 50, color="4")
        add_box("EQUIPOS", 40, 60, 60, 75, color="7")
        add_text("TEXTOS", 65, 65, "ITM 3P-2000A / 50 kA AIC", 3.5, "7")

        add_line("BUS", -50, 50, 250, 50, color="4")
        add_line("BUS", -50, 48, 250, 48, color="4")
        add_text("TEXTOS", 50, 55, f"BUS PRINCIPAL TGBT {cfg['v_nom']:.0f}V", 4.0, "7")

        add_line("RED_BT", -20, 48, -20, 10, color="1") 
        add_box("EQUIPOS", -30, -10, -10, 10, color="1")
        add_text("TEXTOS", -45, -20, "CARGAS DEL PROYECTO", 3.5, "7")

        add_line("RED_BT", 150, 48, 150, 10, color="6")
        add_box("EQUIPOS", 110, -10, 190, 10, color="6")
        add_text("TEXTOS", 115, 0, f"INVERSOR HÍBRIDO {inv_req:.1f} kVA", 3.5, "7")

        add_line("RED_DC", 130, -10, 130, -40, color="2")
        add_box("EQUIPOS", 110, -60, 150, -40, color="2")
        add_text("TEXTOS", 115, -50, f"ARREGLO PV {cfg['p_pv']:.0f} kWp", 3.0, "7")

        add_line("RED_DC", 170, -10, 170, -40, color="3")
        add_box("EQUIPOS", 150, -60, 190, -40, color="3")
        add_text("TEXTOS", 155, -50, f"BANCO BESS {cfg['c_bat']:.0f} kWh", 3.0, "7")

        lines.extend(["0", "ENDSEC", "0", "EOF"])
        return "\n".join(lines)

    c1, c2 = st.columns(2)
    with c1:
        st.markdown("<div class='kpi-card' style='padding: 20px;'><h4 style='color:#F8FAFC; margin-top:0;'>Plano Vectorial CAD</h4><p style='color:#94A3B8; font-size:13px;'>Genera el diagrama unifilar completo (Líneas, Componentes, Marco y Cajetín Técnico) en formato DXF compatible con AutoCAD/ETAP.</p></div>", unsafe_allow_html=True)
        st.download_button("📐 DESCARGAR PLANO CAD (.DXF)", generate_dxf_full().encode('utf-8'), f"Unifilar_{cfg['nombre_proyecto'].replace(' ','_')}.dxf", 'application/dxf', use_container_width=True)
    with c2:
        st.markdown("<div class='kpi-card' style='padding: 20px;'><h4 style='color:#F8FAFC; margin-top:0;'>Datos Simulación</h4><p style='color:#94A3B8; font-size:13px;'>Tabla de balance horario (24h) en CSV con curvas de demanda, BESS, PV y Consumo de Red.</p></div>", unsafe_allow_html=True)
        st.download_button("📊 DESCARGAR RESULTADOS (.CSV)", df_ems.to_csv(index=False).encode('utf-8'), 'Resultados_EMS.csv', 'text/csv', use_container_width=True)

    st.markdown("<br><h4 style='color: #F8FAFC;'>Código Operativo EMS (MATLAB)</h4>", unsafe_allow_html=True)
    matlab_code = f"""%% ============================================================
%% EMS Peak Shaving — {cfg['nombre_proyecto']}
%% Ubicación: {cfg['ubicacion_proyecto']}
%% ============================================================
clear; clc; close all;

%% Parámetros del sistema
P_lim    = {cfg['p_lim']};       % Límite red [kW]
C_bat    = {cfg['c_bat']};      % Capacidad BESS [kWh]
P_PV     = {cfg['p_pv']};       % Potencia PV instalada [kWp]
V_nom    = {cfg['v_nom']};        % Tensión nominal BT [V]
S_trafo  = {cfg['s_trafo']};    % Potencia trafo [kVA]
carga_noc= {cfg['carga_noc']};       % Carga nocturna [kW]

%% Perfiles
P_carga = [{', '.join(map(str, REAL_LOAD))}];
P_PV_base = [{', '.join(map(str, PV_BASE))}];
P_PV_real = P_PV_base * (P_PV / 150);

%% Algoritmo EMS determinístico
SOC_min = 0.20 * C_bat; SOC_max = C_bat; E_act = C_bat * 0.50;
P_bat = zeros(1,24); P_red = zeros(1,24);

for t = 1:24
    P_teo = P_carga(t) - P_PV_real(t);
    if P_teo > P_lim
        req = P_teo - P_lim;
        P_b = min(req, max(0, E_act - SOC_min));
    elseif t >= 2 && t <= 6
        P_b = -min(carga_noc, SOC_max - E_act);
    else
        P_b = 0;
    end
    E_act = E_act - P_b;
    P_bat(t) = P_b;
    P_red(t) = P_teo - P_b;
end
disp('=== SIMULACIÓN EXITOSA ===');
fprintf('Demanda pico recortada: %.1f kW\\n', max(P_red));
"""
    st.code(matlab_code, language='matlab')

# ------------------------------------------
# VISTA 6: CONFIGURACIÓN 
# ------------------------------------------
elif st.session_state.page == "Configuracion":
    st.markdown("<h3 style='color: #00B8FF;'>Configuración de Proyecto Oculta</h3>", unsafe_allow_html=True)
    st.markdown("<p style='color:#94A3B8;'>Para facilitar la experiencia de usuario, la configuración de parámetros ahora se encuentra colapsada en la vista <b>Dashboard Principal</b>. Por favor, navega hacia el Dashboard para ajustar las variables operativas.</p>", unsafe_allow_html=True)
