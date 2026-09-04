import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generar_docx(cfg, inv_req):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1); section.bottom_margin = Inches(1); section.left_margin = Inches(1); section.right_margin = Inches(1)
        
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MEMORIA TÉCNICA Y ESPECIFICACIONES DE PROYECTO"); r.bold = True; r.font.size = Pt(14)
    
    t_meta = doc.add_table(rows=4, cols=2); t_meta.style = 'Table Grid'
    data_m = [("Departamento:", "Ingeniería Técnica"), ("Documento:", f"Memoria Técnica EMS - {cfg['nombre_proyecto']}"), ("Código Documento:", "PROY-EMS-MTC-001"), ("Fecha:", "2026")]
    for i, (k, v) in enumerate(data_m):
        t_meta.cell(i, 0).text = k; t_meta.cell(i, 0).paragraphs[0].runs[0].bold = True; t_meta.cell(i, 1).text = v

    doc.add_heading('1. OBJETIVOS', level=1)
    doc.add_paragraph(f"Electrificación y recorte de demanda pico (Peak Shaving) a {cfg['p_lim']:.1f} kW, garantizando el cumplimiento de normativas de interconexión (IEEE 2800, ARCONEL-001/24).")
    doc.add_heading('2. ESTUDIO TÉCNICO Y ESTABILIDAD', level=1)
    doc.add_paragraph(f"El inversor IBR cumple límites de voltaje continuo (0.9-1.05 p.u.) mediante inyección reactiva dinámica de hasta {inv_req:.1f} kVA.")
    
    target = io.BytesIO()
    doc.save(target)
    return target.getvalue()

def generate_dxf_full(cfg, inv_req):
    lines = ["0", "SECTION", "2", "HEADER", "0", "ENDSEC", "0", "SECTION", "2", "TABLES", "0", "ENDSEC", "0", "SECTION", "2", "BLOCKS", "0", "ENDSEC", "0", "SECTION", "2", "ENTITIES"]
    def add_line(layer, x1, y1, x2, y2, color="7"): lines.extend(["0", "LINE", "8", layer, "62", color, "10", f"{x1:.2f}", "20", f"{y1:.2f}", "30", "0.0", "11", f"{x2:.2f}", "21", f"{y2:.2f}", "31", "0.0"])
    def add_circle(layer, cx, cy, r, color="7"): lines.extend(["0", "CIRCLE", "8", layer, "62", color, "10", f"{cx:.2f}", "20", f"{cy:.2f}", "30", "0.0", "40", f"{r:.2f}"])
    def add_text(layer, x, y, text, height=3.0, color="7"): lines.extend(["0", "TEXT", "8", layer, "62", color, "10", f"{x:.2f}", "20", f"{y:.2f}", "30", "0.0", "40", f"{height:.2f}", "1", str(text)])
    def add_box(layer, x1, y1, x2, y2, color="7"): add_line(layer, x1, y1, x2, y1, color); add_line(layer, x2, y1, x2, y2, color); add_line(layer, x2, y2, x1, y2, color); add_line(layer, x1, y2, x1, y1, color)

    add_box("MARCO", -200, -150, 300, 250, color="2")
    add_box("CAJETIN", 150, -150, 300, -80, color="2")
    add_text("TEXTOS", 160, -90, f"PROYECTO: {cfg['nombre_proyecto'].upper()}", 4.0, "7")
    add_text("TEXTOS", 160, -110, f"UBICACION: {cfg['ubicacion_proyecto'].upper()}", 3.5, "7")
    
    add_line("RED_MT", 50, 220, 50, 160, color="4")
    add_text("TEXTOS", 55, 210, "RED CNEL - 13.8 kV", 3.5, "7")
    add_circle("TRAFO", 50, 115, 15, color="4"); add_circle("TRAFO", 50, 95, 15, color="4")
    add_text("TEXTOS", 85, 120, f"TRAFO {cfg['s_trafo']:.0f} kVA", 3.5, "7")
    
    add_line("BUS", -50, 50, 250, 50, color="4")
    add_text("TEXTOS", 50, 55, f"BUS TGBT {cfg['v_nom']:.0f}V", 4.0, "7")
    
    add_line("RED_BT", -20, 48, -20, 10, color="1"); add_box("EQUIPOS", -30, -10, -10, 10, color="1")
    add_text("TEXTOS", -45, -20, "CARGAS", 3.5, "7")
    
    add_line("RED_BT", 150, 48, 150, 10, color="6"); add_box("EQUIPOS", 110, -10, 190, 10, color="6")
    add_text("TEXTOS", 115, 0, f"INVERSOR {inv_req:.1f} kVA", 3.5, "7")
    
    add_line("RED_DC", 130, -10, 130, -40, color="2"); add_box("EQUIPOS", 110, -60, 150, -40, color="2")
    add_text("TEXTOS", 115, -50, "PV", 3.0, "7")
    add_line("RED_DC", 170, -10, 170, -40, color="3"); add_box("EQUIPOS", 150, -60, 190, -40, color="3")
    add_text("TEXTOS", 155, -50, "BESS", 3.0, "7")

    lines.extend(["0", "ENDSEC", "0", "EOF"])
    return "\n".join(lines)

def generar_codigo_matlab(cfg, kpis):
    matlab_code = f"""%% ============================================================
%% Análisis Dinámico y Peak Shaving EMS — {cfg['nombre_proyecto']}
%% ============================================================
clear; clc;
P_lim = {cfg['p_lim']}; S_inv = {kpis['inv_req']}; V_nom_pu = 1.0;
P_carga = [{', '.join(map(str, kpis['REAL_LOAD']))}];
P_PV_real = [{', '.join(map(str, kpis['pv_real']))}];
P_bat = zeros(1,24); Q_inyectada = zeros(1,24);

for t = 1:24
    P_teo = P_carga(t) - P_PV_real(t);
    P_bat(t) = max(0, P_teo - P_lim);
    
    %% Control Reactivo Volt/VAR (IEEE 2800) Droop Control
    Q_max = sqrt(max(0, S_inv^2 - P_bat(t)^2));
    
    %% Simulación de voltaje
    V_bus = V_nom_pu - (P_teo / ({cfg['s_trafo']} * 0.4)); 
    
    if V_bus < 0.98
        Q_req = (0.98 - V_bus) * (S_inv * 2);
        Q_inyectada(t) = min(Q_req, Q_max);
    elseif V_bus > 1.02
        Q_req = (V_bus - 1.02) * (S_inv * 2);
        Q_inyectada(t) = max(-Q_req, -Q_max);
    end
end
disp('=== SIMULACIÓN COMPLETADA ===');
"""
    return matlab_code
